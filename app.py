import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario de Problemas Filosóficos", page_icon="📚", layout="wide")

# Function to set the background color
def set_background_color(color):
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: {color};
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ### Sobre esta aplicación

    Esta aplicación es un Diccionario de Problemas Filosóficos. Permite a los usuarios obtener respuestas a problemas filosóficos según la interpretación de diversas corrientes filosóficas.

    ### Cómo usar la aplicación:

    1. Elija un problema filosófico de la lista predefinida o proponga su propio problema.
    2. Seleccione una o más corrientes filosóficas.
    3. Haga clic en "Obtener respuesta" para generar las respuestas.
    4. Lea las respuestas y fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 26 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario de Problemas Filosóficos* [Aplicación web]. https://dicproblemasfil.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar respuestas basadas en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario de Problemas Filosóficos")

# Set background color to light yellow
set_background_color("#FFF9C4")  # Light yellow color code

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # List of 101 philosophical problems
    problemas_filosoficos = sorted([
        "¿Qué es la realidad?", "¿Qué es el conocimiento?", "¿Qué es la verdad?", "¿Qué es la conciencia?", 
        "¿Existe el libre albedrío?", "¿Cuál es el propósito de la vida?", "¿Qué es la moral?", 
        "¿Qué es la belleza?", "¿Qué es la justicia?", "¿Existe Dios?", "¿Qué es el tiempo?",
        "¿Qué es el ser?", "¿Qué es la mente?", "¿Qué es la identidad?", "¿Qué son los números?",
        "¿Qué es el lenguaje?", "¿Qué es la ciencia?", "¿Qué es la percepción?", "¿Qué es la felicidad?",
        "¿Qué es el arte?", "¿Qué es el alma?", "¿Qué es la libertad?", "¿Qué es la igualdad?",
        "¿Qué es la virtud?", "¿Qué es el bien?", "¿Qué es el mal?", "¿Qué es el destino?", 
        "¿Qué es el azar?", "¿Qué es la causalidad?", "¿Qué es la necesidad?", "¿Qué es la contingencia?", 
        "¿Qué es la justicia social?", "¿Qué es el poder?", "¿Qué es el deber?", "¿Qué es la responsabilidad?", 
        "¿Qué es el ego?", "¿Qué es el inconsciente?", "¿Qué es la subjetividad?", "¿Qué es la experiencia?",
        "¿Qué es la fenomenología?", "¿Qué es el estructuralismo?", "¿Qué es el deconstruccionismo?", 
        "¿Qué es el nihilismo?", "¿Qué es el relativismo?", "¿Qué es el solipsismo?", "¿Qué es el existencialismo?", 
        "¿Qué es la metafísica?", "¿Qué es la epistemología?", "¿Qué es la estética?", "¿Qué es la política?", 
        "¿Qué es la ética?", "¿Qué es la lógica?", "¿Qué es la dialéctica?", "¿Qué es el materialismo?", 
        "¿Qué es el idealismo?", "¿Qué es el empirismo?", "¿Qué es el racionalismo?", "¿Qué es el pragmatismo?", 
        "¿Qué es el positivismo?", "¿Qué es el constructivismo?", "¿Qué es el naturalismo?", "¿Qué es el humanismo?",
        "¿Qué es la alienación?", "¿Qué es la trascendencia?", "¿Qué es la inmanencia?", "¿Qué es la nada?", 
        "¿Qué es el ser-en-sí?", "¿Qué es el ser-para-sí?", "¿Qué es la esencia?", "¿Qué es la existencia?", 
        "¿Qué es la autenticidad?", "¿Qué es la absurdidad?", "¿Qué es la angustia?", "¿Qué es la desesperación?", 
        "¿Qué es el amor?", "¿Qué es la muerte?", "¿Qué es la eternidad?", "¿Qué es el infinito?",
        "¿Qué es el cosmos?", "¿Qué es la armonía?", "¿Qué es el conflicto?", "¿Qué es la dialéctica?", 
        "¿Qué es la lucha de clases?", "¿Qué es el contrato social?", "¿Qué es la anarquía?", 
        "¿Qué es la utopía?", "¿Qué es la distopía?", "¿Qué es la secularización?", "¿Qué es el dogmatismo?", 
        "¿Qué es el escepticismo?", "¿Qué es el misticismo?", "¿Qué es el panteísmo?", "¿Qué es el dualismo?", 
        "¿Qué es el monismo?", "¿Qué es la pluralidad?", "¿Qué es el universalismo?", "¿Qué es el particularismo?", 
        "¿Qué es el sincretismo?", "¿Qué es el fundamentalismo?", "¿Qué es el multiculturalismo?", "¿Qué es la alteridad?"
    ])

    # List of philosophical schools of thought
    corrientes_filosoficas = [
        "Idealismo", "Realismo", "Existencialismo", "Pragmatismo", "Empirismo", 
        "Racionalismo", "Feminismo", "Positivismo", "Marxismo", "Fenomenología", 
        "Estructuralismo", "Constructivismo", "Posmodernismo", "Nihilismo", "Humanismo"
    ]

    def buscar_informacion(query, corriente):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} {corriente} filosofía"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_respuesta(problema, corriente, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nProblema: {problema}\nCorriente: {corriente}\n\nProporciona una respuesta al problema filosófico '{problema}' según la interpretación del {corriente}. La respuesta debe ser concisa pero informativa, similar a una entrada de diccionario. Si es posible, incluye una referencia a una obra o figura específica de {corriente} que trate este concepto.\n\nRespuesta:",
            "max_tokens": 2048,
            "temperature": 0,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 0,
            "stop": ["Problema:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(problema, respuestas, fuentes):
        doc = Document()
        doc.add_heading('Diccionario de Problemas Filosóficos', 0)

        doc.add_heading('Problema', level=1)
        doc.add_paragraph(problema)

        for corriente, respuesta in respuestas.items():
            doc.add_heading(f'Respuesta según la corriente {corriente}', level=2)
            doc.add_paragraph(respuesta)

        doc.add_heading('Fuentes', level=1)

        # Limitar la lista de fuentes a las primeras 10
        for fuente in fuentes[:10]:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("**Elige un problema filosófico de la lista o propón tu propio problema**:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio problema"])

    if opcion == "Elegir de la lista":
        problema = st.selectbox("Selecciona un problema:", problemas_filosoficos)
    else:
        problema = st.text_input("Ingresa tu propio problema filosófico:")

    st.write("Selecciona una o más corrientes filosóficas (máximo 5):")
    corrientes_seleccionadas = st.multiselect("Corrientes Filosóficas", corrientes_filosoficas)

    if len(corrientes_seleccionadas) > 5:
        st.warning("Has seleccionado más de 5 corrientes. Por favor, selecciona un máximo de 5.")
    else:
        if st.button("Obtener respuesta"):
            if problema and corrientes_seleccionadas:
                with st.spinner("Buscando información y generando respuestas..."):
                    respuestas, todas_fuentes = {}, []

                    for corriente in corrientes_seleccionadas:
                        # Buscar información relevante
                        resultados_busqueda = buscar_informacion(problema, corriente)
                        contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("organic", [])])
                        fuentes = [item["link"] for item in resultados_busqueda.get("organic", [])]

                        # Generar respuesta
                        respuesta = generar_respuesta(problema, corriente, contexto)

                        respuestas[corriente] = respuesta
                        todas_fuentes.extend(fuentes)

                    # Mostrar las respuestas
                    st.subheader(f"Respuestas para el problema: {problema}")
                    for corriente, respuesta in respuestas.items():
                        st.markdown(f"**{corriente}:** {respuesta}")

                    # Botón para descargar el documento
                    doc = create_docx(problema, respuestas, todas_fuentes)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="Descargar respuesta en DOCX",
                        data=buffer,
                        file_name=f"Respuesta_{problema.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Por favor, selecciona un problema y al menos una corriente.")
